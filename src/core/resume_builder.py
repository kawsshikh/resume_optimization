import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import ImageFont
import json

class StyleConfig:
    FONT_NAME = 'Aptos'
    BASE_SIZE = 10
    HEADER_SIZE = 11
    NAME_SIZE = 18
    MARGINS = 0.25
    LINE_SPACING = 1.0
    HYPERLINK_COLOR = "0000FF"
    SECTION_BORDER_SZ = '6'


class DocUtils:
    @staticmethod
    def add_hyperlink(paragraph, url, text, font_name, size):
        part = paragraph.part
        full_url = url if url.startswith('http') else f"https://{url}"
        r_id = part.relate_to(full_url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                              is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        for tag, val in [('w:rStyle', 'Hyperlink'), ('w:sz', str(size * 2)), ('w:color', StyleConfig.HYPERLINK_COLOR),
                         ('w:u', 'single')]:
            el = OxmlElement(tag)
            el.set(qn('w:val'), val)
            rPr.append(el)

        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rPr.append(rFonts)

        new_run.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    @staticmethod
    def add_section_border(paragraph):
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), StyleConfig.SECTION_BORDER_SZ)
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)


class ResumeBuilder:
    def __init__(self, data_path, font_path, reordered_sections):
        with open(data_path, 'r', encoding='utf-8') as f:
            self.data = json.load(f)

        self.doc = Document()
        self.font_path = font_path
        self.font_name = 'Aptos'
        self.header_align = 'left'
        self.font_size = 10
        self.order = reordered_sections

        section = self.doc.sections[0]
        section.top_margin = section.bottom_margin = Inches(0.25)
        section.right_margin = section.left_margin = Inches(0.25)
        self.right_edge = section.page_width - section.left_margin - section.right_margin

    def add_formatted_run(self, paragraph, text, default_bold=False, font_size=10):
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                clean_text = part.replace('**', '')
                run = paragraph.add_run(clean_text)
                run.bold = True
            else:
                run = paragraph.add_run(part)
                run.bold = default_bold
            run.font.name = self.font_name
            run.font.size = Pt(font_size)

    def add_hyperlink(self, url, text):
        part = self.doc.part
        if 'https://' not in url:
            full_address = f"https://{url}"
        else:
            full_address = url
        r_id = part.relate_to(full_address, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'Hyperlink')
        rPr.append(rStyle)

        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '20')
        rPr.append(sz)

        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        rPr.append(color)

        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), self.font_name)
        rFonts.set(qn('w:hAnsi'), self.font_name)
        rPr.append(rFonts)
        new_run.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)
        hyperlink.append(new_run)
        return hyperlink

    def add_element(self, content, paragraph_style="Normal", font_size=10, bold=False, alignment='left'):
        p = self.doc.add_paragraph(style=paragraph_style)
        pf = p.paragraph_format
        if alignment == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == 'right':
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self.add_formatted_run(p, content, default_bold=bold, font_size=font_size)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        pf.space_before = Pt(5 if bold else 0)
        if font_size == 11 and bold:
            self.section_break(p)
        return p

    def styled_element(self, left_text, right_content, tab_pos):
        is_header = (tab_pos == self.right_edge)
        style = None if is_header else 'List Bullet'
        p = self.doc.add_paragraph(style=style)
        pf = p.paragraph_format
        if style == 'List Bullet':
            pf.left_indent = tab_pos
            pf.first_line_indent = -tab_pos
        align = WD_TAB_ALIGNMENT.RIGHT if is_header else WD_TAB_ALIGNMENT.LEFT
        pf.tab_stops.add_tab_stop(tab_pos, alignment=align)
        self.add_formatted_run(p, f"{left_text}\t", default_bold=True, font_size=10)

        if isinstance(right_content, str):
            self.add_formatted_run(p, right_content, default_bold=is_header, font_size=10)
        else:
            p._p.append(right_content)

        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        pf.space_before = Pt(0)
        return p

    def section_break(self, paragraph):
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def get_tab_length(self, skills):
        font_size_px = int(self.font_size * (96 / 72))
        try:
            font = ImageFont.truetype(self.font_path, font_size_px)
        except IOError:
            return 1.5
        max_width_px = 0
        for skill in skills:
            width_px = font.getlength(f"{skill}:")
            max_width_px = max(max_width_px, width_px)
        return (max_width_px / 96) + 0.5

    def personal(self):
        self.add_element(self.data['basics']['name'], font_size=18, bold=True)
        contact = f"{self.data['basics']['location']} | {self.data['basics']['phone']} | {self.data['basics']['email']}"
        self.add_element(contact)
        if self.data['basics'].get("linkedin") or self.data['basics'].get("github"):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            if self.data['basics'].get("linkedin"):
                run = p.add_run("Linkedin: ")
                run.font.size = Pt(10)
                run.font.name = self.font_name
                p._p.append(self.add_hyperlink(self.data['basics']['linkedin'], self.data['basics']['linkedin']))
            if self.data['basics'].get("linkedin") and self.data['basics'].get("github"):
                run = p.add_run(" | ")
                run.font.size = Pt(10)
                run.font.name = self.font_name
            if self.data['basics'].get("github"):
                run = p.add_run("GitHub: ")
                run.font.size = Pt(10)
                run.font.name = self.font_name
                p._p.append(self.add_hyperlink(self.data['basics']['github'], self.data['basics']['github']))

    def summary(self):
        self.add_element("SUMMARY", font_size=11, bold=True)
        self.add_element(self.data["summary"])

    def skills(self):
        self.add_element("SKILLS", font_size=11, bold=True)
        skills_dict = self.data["skills"]
        tab_pos = Inches(self.get_tab_length(list(skills_dict.keys())))
        for key, value in skills_dict.items():
            self.styled_element(key, ', '.join(value), tab_pos)

    def work(self):
        self.add_element("WORK EXPERIENCE", font_size=11, bold=True)
        for work in self.data["professional_experience"]:
            self.styled_element(f"{work['company']} | {work['role']}", work['duration'], self.right_edge)
            for res in work["responsibilities"]:
                self.add_element(res, paragraph_style='List Bullet')

    def projects(self):
        if self.data.get("projects"):
            self.add_element("PROJECTS", font_size=11, bold=True)
            for proj in self.data["projects"]:
                link_el = self.add_hyperlink(proj["link"], "Project Link")
                self.styled_element(proj['title'], link_el, self.right_edge)
                for des in proj["description"]:
                    self.add_element(des, paragraph_style='List Bullet')
                techs = ', '.join(proj["tech_stack"])
                self.add_element(f"**Tech Stack: ** {techs}")

    def Education(self):
        self.add_element("EDUCATION", font_size=11, bold=True)
        for education in self.data["education"]:
            self.add_element(
                f"**{education['degree']}** | {education['institution']}, {education['graduation_date']} | GPA: {education['gpa']}")

    def certifications(self):
        self.add_element("CERTIFICATIONS", font_size=11, bold=True)
        for certificate in self.data["certification"]:
            self.add_element(certificate, paragraph_style='List Bullet')

    def build_resume(self, output_path):
        method_map = {
            "Personal": self.personal,
            "Summary": self.summary,
            "Skills": self.skills,
            "Work Experience": self.work,
            "Education": self.Education,
            "Projects": self.projects,
            "Certification" : self.certifications
        }
        for task in self.order:
            action = method_map.get(task)
            if action:
                action()
        self.doc.save(output_path)